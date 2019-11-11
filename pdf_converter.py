from docx import Document
from lxml import etree

import json
import os

# doc = etree.XMLParser()
doc = etree.XMLParser(recover=True)
document = Document()
top_thresh = 50

# def process_xml_data(input_filepath, output_filepath):
# 	""" Takes the XML input file and saves the data as docx file. """
# 	tree = etree.parse(input_filepath, doc)
# 	pages = tree.findall("page")
# 	page_dimension = get_page_dimension(pages[0])
# 	save_as_docx(pages, output_filepath)

def process_xml_data(input_filepath, output_filepath):
	""" Takes the XML input file and saves the data as docx file. """
	tree = etree.parse(input_filepath, doc)
	pages = tree.findall("page")
	page_dimension = get_page_dimension(pages[0])
	# save_as_docx(pages, output_filepath)
	page_data = get_scores(pages)
	overall_data = {}
	for pg_id, data in page_data.items():
		if not data:
			continue
		ind = get_lowest_top(data)
		print(json.dumps(data, indent = 4))
		print("initial length>> {}".format(len(data)))
		data, new_data = recurssive_data_collection(data, ind, [])
		print("initial new_data length>> {}".format(len(new_data)))
		while data:
			print("Recurssive")
			ind = get_lowest_top(data)
			print("ind>> ", ind)
			data, new_data = recurssive_data_collection(data, ind, new_data)
			print("Recurssive data length>> {}".format(len(data)))
			print("Recurssive new_data length>> {}".format(len(new_data)))
		overall_data[pg_id]	= new_data
	save_as_docx_2(overall_data, output_filepath)

def recurssive_data_collection(data, ind, new_data = []):
	# new_data = []
	try:
		thresh = data[ind][0]
		for inds, each in enumerate(data[ind:]):
			# if (thresh + top_thresh < each[0]):
			if inds == 0 or (abs(each[0] - last_thresh) <= top_thresh):
				pass
			# elif (last_thresh + top_thresh < each[0]):
			else:
				break
			last_index = inds + ind
			last_thresh = each[0]
			last_left = each[1]
			new_data.append(each)
		data = data[:ind] + data[last_index+1:]	
	except:
		print("-")
	return data, new_data

def get_lowest_top(data):
	temp = [i[0] for i in data]
	min_top = min(temp)
	ind = temp.index(min_top)	
	
	for i, each in enumerate(data):
		if abs(min_top - each[0]) <= 5 and each[1] < data[ind][1]:
		# if each[1] < data[ind][1]:
			ind = i
			break 
	
	return ind

def save_as_docx(pages, output_filepath):
	""" Saves the given data as docx formated file. """
	for ind, each in enumerate(pages):
		data = ""
		for i in each.itertext():
			data += i
		document.add_paragraph(data)
		document.add_page_break()
	document.save(output_filepath)

def save_as_docx_2(data, output_filepath):
	""" Saves the given data as docx formated file. """
	# data_to_dump = ""
	last_height = 0
	max_pages = max(data.keys()) + 1

	for i in range(0, max_pages):
		data_to_dump = ""
		for ind, each in enumerate(data.get(i, [])):
			if not each[-1]:
				each[-1] = ""
				continue
			if last_height == each[3]:
				# data_to_dump += " " + each[-1]
				data_to_dump += "\n" + each[-1]
				continue
			elif last_height != each[3]:
				# document.add_paragraph(data_to_dump)		
				data_to_dump += "\n" + each[-1]
			last_height = each[3]
		document.add_paragraph(data_to_dump)
		document.add_page_break()

	# for ind, each in enumerate(data):
	# 	if not each[-1]:
	# 		each[-1] = ""
	# 		continue
	# 	if last_height == each[3]:
	# 		data_to_dump += " " + each[-1]
	# 		continue
	# 	elif last_height != each[3]:
	# 		# document.add_paragraph(data_to_dump)		
	# 		data_to_dump += "\n" + each[-1]
	# 	# else:		
	# 	# data_to_dump = each[-1]
	# 	last_height = each[3]
	# 	# data_to_dump += each[-1]
	# 	print("last_height> ", last_height)
	# document.add_paragraph(data_to_dump)
	# # data_to_dump = ""
	document.add_page_break()
	document.save(output_filepath)

def get_page_dimension(page):
	return page.values()[2:]

def get_approx_footer(page_dimension):
	return page_dimension[0] - top_thresh

def convert_to_pdf_libre(outdir, input_file):
	""" libreoffice --convert-to pdf --outdir outputs outputs/*.docx """
	command = "libreoffice --convert-to pdf --outdir {0} {1}".format(outdir, input_file)

def get_scores(pages):
	font_info = {}
	all_data = {}
	for ind, each_page in enumerate(pages):
		all_data[ind] = []
		for each in each_page:
			temp = []
			if each.tag == "fontspec":
				continue
			temp = [int(i) for i in each.values()[0:4]]
			txt = ""
			for tt in each.itertext():
				txt += tt
			temp.append(txt)	
			# temp.append(each.text)
			all_data[ind].append(temp)
	return all_data

def convert_to_xml(input_filepath, output_filepath):
	""" Takes input_filepath and converts the given pdf file to xml file and saves it in the given output_filepath. """
	try:
		command = "pdftohtml -i -xml {0} {1}".format(input_filepath, output_filepath)
		os.system(command)
	except Exception as e:
		assert("Parsing Error")

if __name__ == "__main__":
	input_filepath = "/home/box/Workshop/training/libin/pdfExtractor/data/Code_of_Conduct.pdf"
	output_filepath = "/home/box/Workshop/training/libin/pdfExtractor/data/Code_of_Conduct.xml"
	convert_to_xml(input_filepath, output_filepath)
	input_filepath = "/home/box/Workshop/training/libin/pdfExtractor/data/Code_of_Conduct.xml"
	output_filepath = "/home/box/Workshop/training/libin/pdfExtractor/data/Code_of_Conduct.docx"
	process_xml_data(input_filepath, output_filepath)
